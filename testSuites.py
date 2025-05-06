import os
import random
import subprocess
from openpyxl import Workbook
from testCase import TestCase
from docx import Document

class TestSuite:
    def __init__(self, test_dir="tests", output_file="testResult.xlsx"):
        """
        Initialize a test suite with:
        - test_dir: Directory to store test files
        - output_file: Excel file to store results
        - algorithms: List of search algorithms to test
        """
        self.test_dir = test_dir
        self.output_file = output_file
        self.algorithms = ["dfs", "bfs", "gbfs", "astar", "iddfs", "beam"]
        self.tests = []

    def generate_tests(self, num_tests):
        """
        Generate a set of test cases with different types:
        - 40% random tests
        - 20% unreachable tests
        - 20% maze tests
        - 20% dense tests
        """
        if not os.path.exists(self.test_dir):
            os.makedirs(self.test_dir)
        
        # Clear previous tests list
        self.tests = []
        self.test_types = {}  # Map filenames to test types
        
        # Distribution of test types
        random_tests = int(num_tests * 0.4)  # 40% random
        unreachable_tests = int(num_tests * 0.2)  # 20% unreachable
        maze_tests = int(num_tests * 0.2)  # 20% maze
        dense_tests = num_tests - random_tests - unreachable_tests - maze_tests  # 20% dense

        test_counter = 0  # Global counter for all tests
        
        # Generate random tests
        for i in range(random_tests):
            rows = random.randint(6, 15)
            cols = random.randint(6, 15)
            goals = random.randint(1, 3)
            walls = random.randint(3, int(rows * cols * 0.2))
            
            test_case = TestCase(rows, cols, goals, walls, "random")
            filename = os.path.join(self.test_dir, f"test{test_counter}.txt")
            test_case.save_to_file(filename)
            self.tests.append(filename)
            self.test_types[filename] = "random"
            test_counter += 1
        
        # Generate unreachable tests
        for i in range(unreachable_tests):
            rows = random.randint(6, 15)
            cols = random.randint(6, 15)
            goals = random.randint(1, 2)
            walls = random.randint(5, 15)
            
            test_case = TestCase(rows, cols, goals, walls, "unreachable")
            filename = os.path.join(self.test_dir, f"test{test_counter}.txt")
            test_case.save_to_file(filename)
            self.tests.append(filename)
            self.test_types[filename] = "unreachable"
            test_counter += 1
        
        # Generate maze tests
        for i in range(maze_tests):
            rows = random.randint(8, 15)
            cols = random.randint(8, 15)
            goals = random.randint(1, 2)
            walls = rows * cols // 3
            
            test_case = TestCase(rows, cols, goals, walls, "maze")
            filename = os.path.join(self.test_dir, f"test{test_counter}.txt")
            test_case.save_to_file(filename)
            self.tests.append(filename)
            self.test_types[filename] = "maze"
            test_counter += 1
        
        # Generate dense tests
        for i in range(dense_tests):
            rows = random.randint(8, 15)
            cols = random.randint(8, 15)
            goals = random.randint(1, 2)
            walls = int(rows * cols * 0.3)
            
            test_case = TestCase(rows, cols, goals, walls, "dense")
            filename = os.path.join(self.test_dir, f"test{test_counter}.txt")
            test_case.save_to_file(filename)
            self.tests.append(filename)
            self.test_types[filename] = "dense"
            test_counter += 1

        print(f"✅ Generated {len(self.tests)} test cases")

    def run_tests(self):
        """
        Run all tests with all algorithms and save results to Excel file.
        For beam search, test with 4 different beam widths.
        """
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "Search Results"
        sheet.append(["Input File", "Test Type", "Algorithm", "Goal Reached", "Nodes Visited", 
                    "Path Length", "Execution Time", "Beam Width", "Memory Used (KB)"])

        total_tests = len(self.tests) * (len(self.algorithms) - 1 + 4)  # 4 beam widths for beam search
        current = 0

        for test_file in self.tests:
            test_type = self.test_types[test_file]
                
            # Run standard algorithms
            for algo in self.algorithms:
                if algo == "beam":
                    # Test beam search with different widths
                    beam_widths = [1, 3, 5, 7]  # Fixed values covering narrow to wider beams
                    for beam_width in beam_widths:
                        current += 1
                        print(f"[{current}/{total_tests}] Running {algo.upper()} (width={beam_width}) on {os.path.basename(test_file)}")
                        result = self._run_algorithm(test_file, algo, beam_width)
                        sheet.append([
                            os.path.basename(test_file),
                            test_type,
                            algo.upper(),
                            result["goal_reached"],
                            result["nodes_visited"],
                            result["path_length"],
                            result["execution_time"],
                            beam_width,
                            result["memory_used"]
                        ])
                else:
                    current += 1
                    print(f"[{current}/{total_tests}] Running {algo.upper()} on {os.path.basename(test_file)}")
                    result = self._run_algorithm(test_file, algo)
                    sheet.append([
                        os.path.basename(test_file),
                        test_type,
                        algo.upper(),
                        result["goal_reached"],
                        result["nodes_visited"],
                        result["path_length"],
                        result["execution_time"],
                        "N/A",
                        result["memory_used"]
                    ])
        
        # Add algorithm complexity analysis
        self.generate_word_report(workbook)

        workbook.save(self.output_file)
        print(f"✅ All tests completed. Results saved to '{self.output_file}'")

    def generate_word_report(self, workbook=None):
        """Generate a detailed Word document report of algorithm performance"""
        # Create a new Word document
        doc = Document()
        
        # Add title
        doc.add_heading('Search Algorithm Performance Analysis', 0)
        
        # Add introduction
        doc.add_paragraph('This report presents a comparative analysis of search algorithms for pathfinding in grid-based environments. The algorithms tested include BFS, DFS, GBFS, A*, IDDFS, and Beam Search with multiple beam widths.')
        
        # Add test summary section
        doc.add_heading('Test Case Summary', level=1)
        
        # Count test types
        test_type_counts = {}
        for test_file, test_type in self.test_types.items():
            if test_type not in test_type_counts:
                test_type_counts[test_type] = 0
            test_type_counts[test_type] += 1
        
        # Create test distribution table
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Test Type'
        hdr_cells[1].text = 'Number of Tests'
        
        # Add rows for each test type
        total_tests = 0
        for test_type in sorted(test_type_counts.keys()):
            row_cells = table.add_row().cells
            row_cells[0].text = test_type.upper()
            row_cells[1].text = str(test_type_counts[test_type])
            total_tests += test_type_counts[test_type]
        
        # Add total row
        row_cells = table.add_row().cells
        row_cells[0].text = 'TOTAL'
        row_cells[1].text = str(total_tests)
        
        # Add explanation of test types
        doc.add_heading('Test Type Descriptions', level=2)
        doc.add_paragraph('• RANDOM: Grids with randomly placed obstacles, creating varied pathfinding challenges.')
        doc.add_paragraph('• MAZE: Complex grids with maze-like obstacle patterns that test algorithm navigation.')
        doc.add_paragraph('• DENSE: Grids with high obstacle density (approximately 30% of cells are walls).')
        doc.add_paragraph('• UNREACHABLE: Grids where goals are deliberately made unreachable to test termination.')
        
        # Add note about test distribution
        doc.add_paragraph(f'A total of {total_tests} test cases were generated with a distribution designed to test algorithms under different conditions. Each algorithm was tested on all grids, with Beam Search additionally tested with multiple width parameters.')
        
        # Load the workbook if not provided
        if workbook is None:
            try:
                import openpyxl
                workbook = openpyxl.load_workbook(self.output_file)
            except:
                doc.add_paragraph("No test results found. Please run tests first.")
                doc.save("Algorithm_Analysis_Report.docx")
                return "Algorithm_Analysis_Report.docx"
        
        # Extract data from workbook
        results_sheet = workbook["Search Results"]
        
        # Same structure as in algorithm_analysis to organize data
        test_types = set(self.test_types.values())
        algorithms = self.algorithms
        
        # Extract performance data
        performance_by_type = {}
        for test_type in test_types:
            performance_by_type[test_type] = {}
            for algo in algorithms:
                performance_by_type[test_type][algo.upper()] = {
                    "runs": 0, 
                    "successes": 0, 
                    "total_time": 0.0,
                    "total_nodes": 0, 
                    "total_path_length": 0,
                    "total_memory": 0.0
                }
        
        # Extract data from results sheet
        for row in list(results_sheet.rows)[1:]:  # Skip header
            test_type = row[1].value
            algo = row[2].value
            goal_reached = row[3].value == "Yes"
            
            try:
                nodes = int(row[4].value) if row[4].value not in ["N/A", None] else 0
                path_length = int(row[5].value) if row[5].value not in ["N/A", None] else 0
                
                # Parse execution time
                time_str = row[6].value
                time_val = 0.0
                if time_str and isinstance(time_str, str):
                    if "ms" in time_str:
                        time_val = float(time_str.replace("ms", ""))
                    elif "s" in time_str:
                        time_val = float(time_str.replace("s", "")) * 1000.0

                # Parse memory usage
                mem_str = row[8].value
                memory_val = 0.0
                if mem_str and isinstance(mem_str, str) and "KB" in mem_str:
                    memory_val = float(mem_str.split()[0])
                
                # Update statistics
                if test_type in performance_by_type and algo in performance_by_type[test_type]:
                    performance_by_type[test_type][algo]["runs"] += 1
                    if goal_reached:
                        performance_by_type[test_type][algo]["successes"] += 1
                        performance_by_type[test_type][algo]["total_time"] += time_val
                        performance_by_type[test_type][algo]["total_nodes"] += nodes
                        performance_by_type[test_type][algo]["total_path_length"] += path_length
                        performance_by_type[test_type][algo]["total_memory"] += memory_val
            except Exception as e:
                pass  # Skip problematic data
        
        # Add overall performance section
        doc.add_heading('Overall Algorithm Performance', level=1)
        
        # Create table for overall performance
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Algorithm'
        hdr_cells[1].text = 'Success Rate'
        hdr_cells[2].text = 'Avg Time (ms)'
        hdr_cells[3].text = 'Avg Nodes'
        hdr_cells[4].text = 'Avg Path Length'
        hdr_cells[5].text = 'Avg Memory (KB)'
        
        # Calculate overall stats
        overall_stats = {}
        for algo in [a.upper() for a in algorithms]:
            overall_stats[algo] = {
                "runs": 0, "successes": 0, "total_time": 0.0,
                "total_nodes": 0, "total_path_length": 0, "total_memory": 0.0
            }
        
        # Aggregate stats across all test types
        for test_type in test_types:
            for algo in [a.upper() for a in algorithms]:
                stats = performance_by_type[test_type][algo]
                overall_stats[algo]["runs"] += stats["runs"]
                overall_stats[algo]["successes"] += stats["successes"]
                overall_stats[algo]["total_time"] += stats["total_time"]
                overall_stats[algo]["total_nodes"] += stats["total_nodes"]
                overall_stats[algo]["total_path_length"] += stats["total_path_length"]
                overall_stats[algo]["total_memory"] += stats["total_memory"]
        
        # Add rows with overall stats
        for algo in sorted([a.upper() for a in algorithms]):
            stats = overall_stats[algo]
            row_cells = table.add_row().cells
            row_cells[0].text = algo
            
            success_rate = f"{(stats['successes'] / stats['runs'] * 100):.1f}%" if stats['runs'] > 0 else "N/A" 
            avg_time = f"{(stats['total_time'] / stats['successes']):.2f}" if stats['successes'] > 0 else "N/A"
            avg_nodes = f"{(stats['total_nodes'] / stats['successes']):.1f}" if stats['successes'] > 0 else "N/A" 
            avg_path = f"{(stats['total_path_length'] / stats['successes']):.1f}" if stats['successes'] > 0 else "N/A"
            avg_memory = f"{(stats['total_memory'] / stats['successes']):.2f}" if stats['successes'] > 0 else "N/A"
            
            row_cells[1].text = success_rate
            row_cells[2].text = avg_time
            row_cells[3].text = avg_nodes
            row_cells[4].text = avg_path
            row_cells[5].text = avg_memory
        
        doc.add_paragraph(f'The table above summarizes the overall performance of each algorithm across all test cases. The success rate is capped at {success_rate} which depends on the number of unreachable tests produced.')
        
        # Add performance by grid type section
        doc.add_heading('Performance by Grid Type', level=1)
        
        # Analyze each grid type
        for test_type in sorted(test_types):
            if test_type == "unreachable":
                continue

            # Add grid type heading
            doc.add_heading(f'{test_type.upper()} Grids', level=2)
            
            # Create table for this grid type
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Algorithm'
            hdr_cells[1].text = 'Success Rate'
            hdr_cells[2].text = 'Avg Time (ms)'
            hdr_cells[3].text = 'Avg Nodes'
            hdr_cells[4].text = 'Avg Path Length'
            hdr_cells[5].text = 'Avg Memory (KB)'
            
            # Add data for each algorithm
            for algo in sorted([a.upper() for a in algorithms]):
                if test_type == "unreachable":
                    continue

                # Get performance data for this algorithm and grid type
                perf = performance_by_type[test_type][algo]
                row_cells = table.add_row().cells
                row_cells[0].text = algo
                
                success_rate = f"{(perf['successes'] / perf['runs'] * 100):.1f}%" if perf['runs'] > 0 else "N/A"
                avg_time = f"{(perf['total_time'] / perf['successes']):.2f}" if perf['successes'] > 0 else "N/A"
                avg_nodes = f"{(perf['total_nodes'] / perf['successes']):.1f}" if perf['successes'] > 0 else "N/A"
                avg_path = f"{(perf['total_path_length'] / perf['successes']):.1f}" if perf['successes'] > 0 else "N/A"
                avg_memory = f"{(perf['total_memory'] / perf['successes']):.2f}" if perf['successes'] > 0 else "N/A"
                
                row_cells[1].text = success_rate
                row_cells[2].text = avg_time
                row_cells[3].text = avg_nodes
                row_cells[4].text = avg_path
                row_cells[5].text = avg_memory
            
            # Add recommendations for this grid type
            if test_type != "unreachable":
                doc.add_heading(f'Recommendations for {test_type.upper()} Grids', level=3)
                
                # Find best algorithms
                best_overall = self._find_best_overall(performance_by_type[test_type])
                fastest = self._find_fastest(performance_by_type[test_type])
                best_memory = self._find_best_memory(performance_by_type[test_type])
                
                # Add recommendation paragraph
                doc.add_paragraph(f'• Best Overall Algorithm: {best_overall}')
                doc.add_paragraph(f'• Fastest Algorithm: {fastest}')
                doc.add_paragraph(f'• Most Memory Efficient: {best_memory}')
        
        # Add beam search analysis
        doc.add_heading('Beam Search Analysis', level=1)

        # Extract beam search data by beam width
        beam_data = {}
        beam_widths = []

        # Parse results for beam search performance data
        for row in list(results_sheet.rows)[1:]:  # Skip header
            if row[2].value == "BEAM":
                width = row[7].value
                if width not in beam_widths:
                    beam_widths.append(width)
                    beam_data[width] = {
                        "runs": 0,
                        "successes": 0,
                        "total_time": 0,
                        "total_nodes": 0,
                        "total_path": 0,
                        "total_memory": 0
                    }
                
                # Update statistics
                beam_data[width]["runs"] += 1
                goal_reached = row[3].value == "Yes"
                
                if goal_reached:
                    beam_data[width]["successes"] += 1
                    
                    # Extract metrics for successful runs
                    try:
                        nodes = int(row[4].value) if row[4].value not in ["N/A", None] else 0
                        path_length = int(row[5].value) if row[5].value not in ["N/A", None] else 0
                        
                        # Time
                        time_str = row[6].value
                        time_val = 0.0
                        if time_str and isinstance(time_str, str):
                            if "ms" in time_str:
                                time_val = float(time_str.replace("ms", ""))
                            elif "s" in time_str:
                                time_val = float(time_str.replace("s", "")) * 1000.0
                        
                        # Memory
                        mem_str = row[8].value
                        memory_val = 0.0
                        if mem_str and isinstance(mem_str, str) and "KB" in mem_str:
                            memory_val = float(mem_str.split()[0])
                            
                        # Update totals
                        beam_data[width]["total_time"] += time_val
                        beam_data[width]["total_nodes"] += nodes
                        beam_data[width]["total_path"] += path_length
                        beam_data[width]["total_memory"] += memory_val
                        
                    except Exception as e:
                        pass  # Skip problematic data

        # Add introduction to beam search
        doc.add_paragraph('Beam Search is a heuristic search algorithm that limits the breadth of the search by keeping only the k most promising nodes at each level, where k is the beam width. This analysis examines how different beam widths affect performance across grid types.')

        # Create beam width comparison table
        doc.add_heading('Beam Width Performance Comparison', level=2)
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'

        # Add headers
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Beam Width'
        hdr_cells[1].text = 'Success Rate'
        hdr_cells[2].text = 'Avg Time (ms)'
        hdr_cells[3].text = 'Avg Nodes'
        hdr_cells[4].text = 'Avg Path Length'
        hdr_cells[5].text = 'Avg Memory (KB)'

        # Add a row for each beam width
        for width in sorted(beam_widths):
            stats = beam_data[width]
            row_cells = table.add_row().cells
            row_cells[0].text = str(width)
            
            # Calculate averages and success rate
            success_rate = f"{(stats['successes'] / stats['runs'] * 100):.1f}%" if stats['runs'] > 0 else "N/A"
            avg_time = f"{(stats['total_time'] / stats['successes']):.2f}" if stats['successes'] > 0 else "N/A"
            avg_nodes = f"{(stats['total_nodes'] / stats['successes']):.1f}" if stats['successes'] > 0 else "N/A"
            avg_path = f"{(stats['total_path'] / stats['successes']):.1f}" if stats['successes'] > 0 else "N/A"
            avg_memory = f"{(stats['total_memory'] / stats['successes']):.2f}" if stats['successes'] > 0 else "N/A"
            
            # Add to table
            row_cells[1].text = success_rate
            row_cells[2].text = avg_time
            row_cells[3].text = avg_nodes
            row_cells[4].text = avg_path
            row_cells[5].text = avg_memory

        # Add beam search recommendations
        doc.add_heading('Beam Search Recommendations', level=2)

        # Find optimal beam width based on success rate and efficiency
        optimal_width = None
        best_score = float('-inf')

        for width in beam_widths:
            if beam_data[width]["runs"] > 0 and beam_data[width]["successes"] > 0: # Check if the width has been tested
                success_rate = (beam_data[width]["successes"] / beam_data[width]["runs"]) * 100
                avg_time = beam_data[width]["total_time"] / beam_data[width]["successes"]
                avg_memory = beam_data[width]["total_memory"] / beam_data[width]["successes"]
                
                # Score balances success rate with time and memory efficiency
                score = success_rate - (avg_time / 50) - (avg_memory / 100)
                
                if score > best_score: # Update best score
                    best_score = score
                    optimal_width = width

        if optimal_width:
            doc.add_paragraph(f'• Recommended beam width: {optimal_width}')
            doc.add_paragraph(f'• Beam search with width={optimal_width} achieves the best balance of success rate, execution time, and memory usage.')

        # Add observations about beam search behavior
        doc.add_heading('Key Observations', level=3)
        doc.add_paragraph('• Increasing beam width generally improves success rate but increases memory usage and computation time.')
        doc.add_paragraph('• Beam search offers a middle ground between the completeness of BFS and the efficiency of greedy search algorithms.')
        doc.add_paragraph('• For complex grids (dense, maze), larger beam widths are typically more effective.')
        doc.add_paragraph('• The optimal beam width depends on the specific grid complexity and search requirements.')
        
        # Save the report
        report_filename = "Search_Algorithm_Analysis_Report.docx"
        doc.save(report_filename)
        print(f"✅ Word report generated: {report_filename}")
        
        return report_filename

    def _run_algorithm(self, test_file, algorithm, beam_width=None):    
        """
        Run a single algorithm on a test file and return results:
        - goal_reached: Whether the goal was found
        - nodes_visited: Number of nodes explored
        - path_length: Length of the path found
        - execution_time: Time taken to run the algorithm
        """
        try:
            cmd = ["python", "search.py", test_file, algorithm]
            if beam_width and algorithm == "beam":
                cmd.append(str(beam_width))

            process = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, timeout=30)
            output = process.stdout

            # Parse output
            goal_reached = "Goal reached:" in output
            
            # Extract nodes visited
            nodes_visited = "0"
            for line in output.split('\n'):
                if "Nodes visited:" in line:
                    try:
                        nodes_visited = line.split("Nodes visited:")[1].strip()
                    except:
                        pass
            
            # Extract execution time
            execution_time = 0
            for line in output.split('\n'):
                if "Execution time:" in line:
                    try:
                        time_str = line.split("Execution time:")[1].strip()
                        execution_time = float(time_str.split()[0])
                    except:
                        pass
            
            # Extract memory used
            memory_used = 0
            for line in output.split('\n'):
                if "Memory used:" in line:
                    try:
                        memory_str = line.split("Memory used:")[1].strip()
                        memory_used = float(memory_str.split()[0])
                    except:
                        pass
            
            # Extract path length if goal reached
            path_length = 0
            if goal_reached:
                for line in output.split('\n'):
                    if "Path:" in line:
                        try:
                            path_str = line.split("Path:")[1].strip()
                            path_length = len(path_str.split())
                        except:
                            pass
            
            return {
                "goal_reached": "Yes" if goal_reached else "No",
                "nodes_visited": nodes_visited,
                "path_length": path_length,
                "execution_time": f"{execution_time:.3f}ms",
                "memory_used": f"{memory_used:.2f} KB"
            }
            
        except subprocess.TimeoutExpired:
            return {
                "goal_reached": "Timeout",
                "nodes_visited": "N/A",
                "path_length": "N/A",
                "execution_time": "30s+",
                "memory_used": "N/A"
            }
        except Exception as e:
            return {
                "goal_reached": "Error: " + str(e),
                "nodes_visited": "N/A",
                "path_length": "N/A",
                "execution_time": "Error: " + str(e),
                "memory_used": "N/A"
            }

    def _find_best_overall(self, algorithms_perf):
        """Find the best overall algorithm balancing success rate, speed, and memory"""
        candidates = []
        
        for algo, perf in algorithms_perf.items():
            if perf.get("runs", 0) > 0 and perf.get("successes", 0): # Check if the algorithm has been run
                # Calculate score based on success rate, time, and memory
                success_rate = perf.get("successes", 0) / perf.get("runs", 1) * 100
                
                # Only calculate time and memory if there were successes
                if perf.get("successes", 0) > 0:
                    avg_time = perf.get("total_time", 0) / perf.get("successes", 1) # Calculate average time
                    avg_memory = perf.get("total_memory", 0) / perf.get("successes", 1) # Calculate average memory
                    
                    # Lower is better - normalize and weight the score components
                    score = (
                        (100 - success_rate) * 0.4 +  # Success rate (40% weight)
                        (min(avg_time / 100, 1.0)) * 0.3 +  # Time (30% weight)
                        (min(avg_memory / 1000, 1.0)) * 0.3  # Memory (30% weight)
                    )
                    
                    candidates.append((algo, score))
        
        # Return the algorithm with the best (lowest) score
        if candidates:
            best_score = float('inf') # Initialize to infinity
            best_algo = None

            for algo, score in candidates:
                if score < best_score: # Update if this algorithm has a better score
                    best_score = score
                    best_algo = algo
            
            return best_algo 

    def _find_fastest(self, algorithms_perf):
        """Find the fastest algorithm among those with good success rates"""
        best_time = float('inf') # Initialize to infinity
        best_algo = None
        
        for algo, perf in algorithms_perf.items(): 
            if perf.get("runs", 0) > 0 and perf.get("successes", 0) > 0: # Check if the algorithm has been run
                avg_time = perf.get("total_time", 0) / perf.get("successes", 1) # Calculate average time
                if avg_time < best_time: # Update if this algorithm is faster
                    best_time = avg_time
                    best_algo = algo
        
        return best_algo

    def _find_best_memory(self, algorithms_perf):
        """Find the algorithm with lowest memory usage among those with good success rates"""
        best_memory = float('inf') # Initialize to infinity
        best_algo = None
        
        for algo, perf in algorithms_perf.items():
            if perf.get("runs", 0) > 0 and perf.get("successes", 0) > 0: # Check if the algorithm has been run
                avg_memory = perf.get("total_memory", 0) / perf.get("successes", 1) # Calculate average memory
                if avg_memory < best_memory: # Update if this algorithm uses less memory
                    best_memory = avg_memory
                    best_algo = algo
        
        return best_algo

if __name__ == "__main__":
    # Create and run a test suite with 10 test cases
    suite = TestSuite()
    suite.generate_tests(10)
    suite.run_tests()